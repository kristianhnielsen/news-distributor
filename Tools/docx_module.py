import docx
from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor
import datetime


def edit_docx_style(paragragh_object, document_object, existing_style_name, font_size, is_bold):
    paragragh_object.style = document_object.styles[existing_style_name]
    paragragh_object.style.font.name = 'Calibri'
    paragragh_object.style.font.size = Pt(font_size)
    paragragh_object.style.font.color.rgb = RGBColor(0, 0, 0)
    paragragh_object.style.font.bold = is_bold
    paragragh_object.paragraph_format.space_before = Pt(0)
    paragragh_object.paragraph_format.space_after = Pt(0)
    paragragh_object.paragraph_format.line_spacing = 1

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

def create_docx_style(document_object, style_name, based_on, use_font_name, font_size, is_bold, r, g, b):
    styles = document_object.styles
    text_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    text_style.base_style = styles[based_on]
    text_style.font.name = use_font_name
    text_style.font.size = Pt(font_size)
    text_style.font.bold = is_bold
    text_style.font.color.rgb = RGBColor(r, g, b)

def convert_vault(database):
    """ Takes teh vault database, sorts it according to media name and publishing date (newest first), and deletes duplicates. """
    # not good code, but it works
    no_textbody_duplicates_db = list({v['text body']:v for v in database}.values())
    pub_date_sorted_db = sorted(no_textbody_duplicates_db, key = lambda i: i['publishing date'], reverse=True)
    media_sorted_db = sorted(pub_date_sorted_db, key = lambda i: i['media'])
    
    return media_sorted_db

def write_to_docx(database: list, document_name):
    """ Takes an unsorted database, and sorts it, then runs each key in every article into a .docx file.
    document_name has to contain .docx-extension. """
    document_directory = f"{document_name}"
    doc = Document()
    sorted_database = convert_vault(database)
    
    # Add robot disclaimer
    disclaimer = f'This document was created as part of an automated process on {datetime.date.today()}.\nMake sure to open the "Headings"-section of the Navigation Pane (can be found in the "View"-section of the document)'
    doc.add_paragraph(disclaimer)

    # Add name of media as header
    previous_article_media_source = ''
    for article in sorted_database:
        if article['media'] != previous_article_media_source:
            # Add the name of the media as header
            media_title = doc.add_paragraph(article['media'])
            edit_docx_style(media_title, doc, 'Heading 1', 16, True)

        # Add headline
        headline_paragraph = doc.add_paragraph(article['headline'])
        edit_docx_style(headline_paragraph, doc, 'Heading 2', 13, True)

        # Add body, publishing date and source
        body_paragraph = doc.add_paragraph(article['publishing date'] + '\n\n' + article['text body'] + '\n\nSource: \n')
        add_hyperlink(body_paragraph, article['source'], article['source'])
        doc.add_paragraph('\n\n')
        edit_docx_style(body_paragraph, doc, 'Normal', 12, False)

        previous_article_media_source = article['media']
    try:
        doc.save(document_directory)
    except PermissionError:
        print("Access to the file denied. Make sure the file is not already in use!")
        return None


